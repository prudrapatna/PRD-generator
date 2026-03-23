from docx import Document

def check_and_add_warning(file_path):
    doc = Document(file_path)
    warning_text = "Not intended to detect breathing emergences at elevation of approximately 2,500 meters (8,200 feet) or higher"
    
    found = False
    for p in doc.paragraphs:
        if warning_text in p.text:
            found = True
            break
            
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    if warning_text in p.text:
                        found = True
                        break
                        
    if not found:
        print(f"Warning missing in {file_path}. Adding it to the end.")
        doc.add_paragraph(warning_text)
        doc.save(file_path)
    else:
        print(f"Warning already present in {file_path}.")

check_and_add_warning("/Users/prudrapatna/.gemini/tmp/prd-generator/doc1_modified.docx")
check_and_add_warning("/Users/prudrapatna/.gemini/tmp/prd-generator/doc3_modified.docx")
