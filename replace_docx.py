import sys
from docx import Document

def replace_text_in_doc(file_path):
    doc = Document(file_path)
    replacements = {
        "93%": "92%",
        "Impaired Oxygenation Detection": "Breathing Emergency Detection",
        "impaired oxygenation detection": "breathing emergency detection",
        "Impaired oxygenation": "Breathing emergency",
        "impaired oxygenation": "breathing emergency",
        "Monitoring": "Detection",
        "monitoring": "detection",
        "Alarm": "Alert",
        "alarm": "alert",
        "Congestive Heart Failure": "Heart Failure",
        "CHF": "Heart Failure",
        "Adults over age 18": "individuals over age 18",
        "adults over age 18": "individuals over age 18",
        "Continuous": "Passive",
        "continuous": "passive"
    }

    def replace_in_text(text):
        for k, v in replacements.items():
            text = text.replace(k, v)
        return text

    # Process paragraphs
    for p in doc.paragraphs:
        for run in p.runs:
            if run.text:
                run.text = replace_in_text(run.text)

    # Process tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    for run in p.runs:
                        if run.text:
                            run.text = replace_in_text(run.text)

    doc.save(file_path.replace(".docx", "_modified.docx"))
    print(f"Processed {file_path}")

replace_text_in_doc("/Users/prudrapatna/.gemini/tmp/prd-generator/doc1.docx")
replace_text_in_doc("/Users/prudrapatna/.gemini/tmp/prd-generator/doc3.docx")
